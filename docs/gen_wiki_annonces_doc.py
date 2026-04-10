#!/usr/bin/env python3
"""
Génère le document Word : Base de connaissances & Annonces — Documentation complète
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Marges ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

CARE_GREEN = RGBColor(0x2D, 0x4A, 0x43)
ACCENT_ORANGE = RGBColor(0xC5, 0x4A, 0x3A)
GREY_TEXT = RGBColor(0x6B, 0x6B, 0x6B)

for level, size in [(1, 20), (2, 16), (3, 13), (4, 11)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = CARE_GREEN
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.bold = True

# ── Helpers ──
def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if italic: r.italic = True
    if bold: r.bold = True
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(' — ' + text)
    else:
        p.add_run(text)
    return p

def info_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('💡 ' + text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY_TEXT

def code_inline(p, text):
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    return r

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════
#  TITRE
# ═══════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Base de connaissances & Annonces officielles')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = CARE_GREEN

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Documentation fonctionnelle complète — SpocSpace / SpocCare')
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY_TEXT

doc.add_paragraph()
sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sep.add_run('— Phase 1 + 2 + 3 —')
r.font.size = Pt(11)
r.font.color.rgb = ACCENT_ORANGE

doc.add_paragraph()
intro = doc.add_paragraph()
r = intro.add_run('Document de référence sur le module Wiki + Annonces de SpocSpace : ce qui existe, pourquoi ces choix, ce qui peut être ajouté ou modifié.')
r.italic = True
r.font.color.rgb = GREY_TEXT

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  1. VISION
# ═══════════════════════════════════════════════════════
h1('1. Vision et positionnement')

h2('1.1 Le problème métier')
para("Dans un EMS, l'information critique (protocoles de soins, procédures d'urgence, consignes RH, menus) est dispersée : classeurs papier, mails, post-its, dossier partagé. Conséquences :")
bullet("Les soignants ne trouvent pas l'information au bon moment")
bullet("Les versions obsolètes circulent")
bullet("Les nouveaux collaborateurs perdent du temps en formation")
bullet("Aucun moyen de prouver que tel protocole a été lu par l'équipe")
bullet("Aucune visibilité sur ce qui manque dans la base")

h2('1.2 Notre réponse')
para("Deux outils complémentaires intégrés à SpocCare :")
add_table(
    ['Outil', 'Fonction', 'Inspiration'],
    [
        ['Base de connaissances (Wiki)', 'Stockage durable, recherche, structuré', 'Notion / Confluence / Guru'],
        ['Annonces officielles', 'Communication descendante one-shot', 'SharePoint News / Slack #announcements'],
    ]
)
para("Différence fondamentale :", bold=True)
bullet("Le Wiki est de l'information durable qu'on consulte quand on en a besoin")
bullet("Une Annonce est de l'information ponctuelle qu'on pousse vers les gens")

# ═══════════════════════════════════════════════════════
#  2. ARCHITECTURE
# ═══════════════════════════════════════════════════════
h1('2. Architecture technique')

h2('2.1 Stockage')
bullet("MySQL en utf8mb4_unicode_ci")
bullet("11 tables wiki : wiki_pages, wiki_categories, wiki_tags, wiki_page_tags, wiki_versions, wiki_favoris, wiki_page_permissions, wiki_page_views, wiki_search_log, wiki_page_reviews, wiki_suggestions_log")
bullet("3 tables annonces : annonces, annonce_views, annonce_acks")

h2('2.2 Endpoints API')
bullet("Admin / SpocCare", "/spocspace/admin/api.php → admin_api_modules/wiki.php + annonces.php")
bullet("Employé SPA", "/spocspace/api.php → api_modules/wiki.php")
bullet("Recherche globale unifiée", "admin_global_search, admin_care_global_search, global_search")

h2('2.3 Pages utilisateur')
bullet("/spoccare/wiki", "liste + lecture")
bullet("/spoccare/wiki-edit", "éditeur (responsables uniquement)")
bullet("/spoccare/annonces", "liste + lecture")
bullet("/spoccare/annonce-edit", "éditeur")
bullet("/spocspace/admin/wiki", "wrapper d'inclusion")
bullet("/spocspace/admin/annonces", "wrapper d'inclusion")
bullet("/spocspace/admin/wiki-analytics", "dashboard responsables")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  3. WIKI
# ═══════════════════════════════════════════════════════
h1('3. Fonctionnalités du Wiki')

h2('3.1 Catégories')
para("Quoi :", bold=True)
para("Regroupement thématique des pages (ex : Hygiène, Soins infirmiers, RH, Cuisine).")
para("Pourquoi :", bold=True)
bullet("Navigation par thème quand on ne sait pas exactement ce qu'on cherche")
bullet("Couleur + icône → repérage visuel rapide")
bullet("Crée un sentiment d'ordre dans la base")
para("Alternatives écartées :", bold=True)
bullet("Catégories hiérarchiques", "trop complexe pour un EMS, on a privilégié catégories plates + tags")
bullet("Auto-catégorisation IA", "pas fiable, le manuel reste plus simple")

h2('3.2 Tags')
para("Quoi :", bold=True)
para("Étiquettes transversales attachées à une page (ex : #urgent, #nuit, #protocole).")
para("Différence catégorie vs tag :", bold=True)
bullet("Catégorie", "1 et 1 seule (où la page « vit »)")
bullet("Tag", "0 ou plusieurs (comment elle est trouvée)")
info_box("Une page peut avoir plusieurs facettes — un protocole peut être à la fois #hygiène et #urgence")

h2('3.3 Recherche FULLTEXT')
para("Pourquoi le FULLTEXT plutôt que LIKE ?", bold=True)
bullet("LIKE %mot% est lent sur de gros contenus et ne classe pas par pertinence")
bullet("FULLTEXT MySQL est indexé → réponse en quelques ms même sur 10 000 pages")
bullet("Permet la recherche multi-mots, exclusions (-mot), exigences (+mot)")
para("Index : FULLTEXT(titre, description, contenu) créé en migration 058_wiki_phase1.sql")
info_box("Logging automatique de la query et du nombre de résultats — alimente les Knowledge Gaps")

h2('3.4 Favoris personnels')
para("Quoi :", bold=True)
para("Chaque utilisateur peut épingler des pages dans ses favoris (cœur).")
para("Pourquoi :", bold=True)
bullet("Une infirmière de nuit a besoin d'un accès rapide à 5-10 protocoles spécifiques")
bullet("Évite de re-chercher les mêmes pages tous les jours")
bullet("Filtre « Mes favoris » en haut de la liste")

h2('3.5 Verification (cycle de vérification)')
para("Quoi :", bold=True)
para("Chaque page peut avoir un expert assigné (responsable, IDE, médecin) qui doit la revérifier régulièrement (par défaut tous les 90 jours).")
para("Pourquoi :", bold=True)
bullet("Dans la santé, une procédure obsolète peut être dangereuse")
bullet("Force la maintenance proactive (pas seulement réactive)")
bullet("Donne une marque de confiance visible : « Vérifié par X le ... »")
bullet("Inspiré directement de Guru « Knowledge Verification »")
para("Logique :", bold=True)
bullet("À l'assignation : on définit expert_id + verify_interval_days")
bullet("À chaque vérification : verified_at = NOW(), verify_next = NOW() + interval")
bullet("Quand verify_next ≤ NOW() : badge rouge « À revérifier » + bandeau d'alerte")

h2('3.6 Permissions par rôle')
para("Quoi :", bold=True)
para("Une page peut être restreinte à certains rôles (ex : « RH » visible uniquement par Direction + Admin).")
para("Logique inclusive :", bold=True)
para("Si AUCUNE permission n'est définie, la page est visible par tous. Sinon, seuls les rôles listés voient la page.")
info_box("Le rôle est plus simple à gérer que les permissions par utilisateur — adapté à un EMS de 50-200 personnes")

h2('3.7 Versionnage')
para("Quoi :", bold=True)
para("Chaque modification du contenu sauvegarde l'ancienne version dans wiki_versions.")
para("Pourquoi :", bold=True)
bullet("Permet de revenir en arrière en cas d'erreur")
bullet("Trace qui a modifié quoi et quand")
bullet("Audit légal (santé : on doit pouvoir prouver la version en vigueur à une date X)")

h2('3.8 Suggestions IA')
para("Quoi :", bold=True)
para("Sur la page d'accueil du wiki, propose des pages contextuellement pertinentes.")
para("Implémentation actuelle (basique) :", bold=True)
bullet("Algorithme heuristique : pages récentes + tags fréquents + page la plus consultée")
bullet("Table wiki_suggestions_log pour ne pas re-suggérer les mêmes")
para("Évolutions possibles :", bold=True)
bullet("Brancher une vraie IA (Claude, GPT)")
bullet("Apprentissage des clics utilisateur")

h2('3.9 Image de couverture')
para("Bandeau image en haut de chaque page (comme les articles de blog).")
bullet("Modal de choix : Upload local OU recherche Pixabay (banque libre de droits)")
bullet("Conversion automatique en webp pour optimisation")
bullet("Stockage public dans /spocspace/assets/uploads/wiki/")

h2('3.10 Import Word / PDF')
para("Pourquoi :", bold=True)
bullet("95% des EMS ont déjà leurs procédures en Word ou PDF")
bullet("Sans import, la migration est bloquante (personne ne va re-saisir 200 pages)")
bullet("Permet une adoption rapide")
para("Implémentation :", bold=True)
bullet("mammoth.js", "pour les .docx (préserve titres, listes, gras)")
bullet("pdf.js", "pour les PDF (extraction texte)")

h2('3.11 Éditeur TipTap')
para("Pourquoi TipTap ?", bold=True)
bullet("Open source, gratuit, moderne")
bullet("Architecture extensible (ProseMirror)")
bullet("Sortie HTML propre")
bullet("Même rendu en édition et en lecture")
para("Toolbar disponible :", bold=True)
bullet("Bold, Italic, Underline, Strike, Highlight")
bullet("Titres H2/H3, listes, blockquote")
bullet("Alignement (gauche/centre/droite/justifié)")
bullet("Liens (avec modal personnalisé)")
bullet("Images (avec resize 25/50/75/100%)")
bullet("Tableaux (insertion, ajout/suppression de lignes/colonnes)")
bullet("Emoji picker, Undo / Redo")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  4. WORKFLOW
# ═══════════════════════════════════════════════════════
h1('4. Workflow de publication (Phase 3)')

h2('4.1 Les 3 statuts')
add_table(
    ['Statut', 'Visible par', 'Quand l\'utiliser'],
    [
        ['Brouillon', 'Auteur seulement', 'Travail en cours'],
        ['En review', 'Responsables', 'Soumis pour validation'],
        ['Publié', 'Tous (selon permissions)', 'Validé, en production'],
    ]
)

h2('4.2 Pourquoi ce workflow ?')
bullet("Évite la publication directe d'infos non validées")
bullet("Protocoles de soins : un brouillon peut contenir des erreurs dangereuses")
bullet("Permet la collaboration : un junior rédige, un senior valide")
bullet("Inspiré de Confluence « Draft → Review → Published »")

h2('4.3 Actions de review')
bullet("Approuver", "passe en publié automatiquement")
bullet("Demander des modifications", "repasse en brouillon avec un commentaire")
bullet("Commenter", "reste en review, ajoute juste un commentaire au log")

h2('4.4 Historique des reviews')
para("Table wiki_page_reviews :")
bullet("Toutes les décisions sont conservées (qui, quand, quoi, pourquoi)")
bullet("Affichage chronologique sous l'éditeur")
bullet("Audit complet du processus")

h2('4.5 Évolutions possibles')
bullet("Reviewer désigné", "forcer une personne spécifique à valider")
bullet("Multi-review", "exiger N approbations avant publication")
bullet("Notification", "email / push à l'auteur quand sa page est approuvée/rejetée")
bullet("Diff visuel", "montrer en rouge/vert ce qui a changé entre 2 versions")

# ═══════════════════════════════════════════════════════
#  5. ANALYTICS
# ═══════════════════════════════════════════════════════
h1('5. Analytics (Phase 3)')

h2('5.1 Page Vues')
para("Chaque ouverture d'une page wiki est loggée dans wiki_page_views (page_id, user_id, viewed_at).")
para("Pourquoi :", bold=True)
bullet("Mesurer ce qui est réellement utilisé")
bullet("Identifier les pages stars")
bullet("Identifier les pages mortes")
para("KPIs affichés :", bold=True)
bullet("Total pages")
bullet("Vues sur la période (7/30/90 jours)")
bullet("Lecteurs uniques")
bullet("Vérifications expirées")

h2('5.2 Top 10 pages')
para("Classement des pages par nombre de vues sur la période.")
bullet("Mettre en avant les contenus qui marchent")
bullet("Comprendre les besoins de l'équipe")
bullet("Justifier l'effort éditorial")

h2('5.3 Cartes orphelines')
para("Définition exacte :", bold=True)
bullet("Page active (non archivée)")
bullet("Pas modifiée depuis 90+ jours")
bullet("0 vue sur la période sélectionnée")
para("Pourquoi ces critères ?", bold=True)
bullet("Une page sans màj récente ET sans lecture est probablement obsolète ou inutile")
bullet("Si elle est juste sans màj mais lue → c'est un classique stable, on ne touche pas")
bullet("Si elle est juste sans lecture mais récente → on lui laisse sa chance")
info_box("Action recommandée : archiver, fusionner ailleurs, ou retravailler le SEO interne (tags, catégorie)")

h2('5.4 Évolutions possibles')
bullet("Sparkline : courbe des vues jour par jour")
bullet("Heatmap : heure de la journée / jour de la semaine")
bullet("Temps de lecture moyen (nécessite un beacon JS)")
bullet("Bounce rate : % de gens qui ouvrent et ferment immédiatement")
bullet("Export CSV des analytics")
bullet("Comparaison période A vs B")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  6. KNOWLEDGE GAPS
# ═══════════════════════════════════════════════════════
h1('6. Knowledge Gaps (Phase 3)')

h2('6.1 Qu\'est-ce qu\'un gap ?')
para("Une recherche utilisateur qui ne donne aucun résultat. C'est la preuve qu'il manque une page dans la base.")

h2('6.2 Le mécanisme de log')
para("À chaque recherche dans le wiki, on appelle :")
p = doc.add_paragraph()
code_inline(p, "adminApiPost('admin_log_wiki_search', { q: 'plaies de pression', results_count: 0 })")
para("Stocké dans wiki_search_log (id, user_id, q, results_count, created_at).")

h2('6.3 Affichage admin')
para("Dans /spocspace/admin/wiki-analytics → section « Knowledge Gaps » :")
bullet("Liste triée par fréquence (la query la plus cherchée en premier)")
bullet("Badge avec le nombre de fois où elle a été tapée")
bullet("Bouton « Créer » qui ouvre l'éditeur wiki avec le titre pré-rempli")

h2('6.4 Pourquoi c\'est précieux')
bullet("C'est une liste de courses prioritaire pour l'équipe éditoriale")
bullet("Mesurable : « On a comblé 80% des gaps en 3 mois »")
bullet("Aligne le contenu sur les vrais besoins de l'équipe (pas ce que la direction pense)")

h2('6.5 Évolutions possibles')
bullet("Suggestions automatiques d'IA", "« Cette query revient 9× — voici un brouillon Claude à valider »")
bullet("Notification", "alerter le responsable wiki quand une nouvelle query orpheline atteint X occurrences")
bullet("Regrouper les variantes", "« plaie pression », « plaies de pression », « ulcère de pression » → 1 entrée")
bullet("Détection de typos", "« désinfction » → suggérer « désinfection »")

# ═══════════════════════════════════════════════════════
#  7. ANNONCES
# ═══════════════════════════════════════════════════════
h1('7. Annonces officielles')

h2('7.1 Différence avec le mur social')
add_table(
    ['Critère', 'Mur social', 'Annonces officielles'],
    [
        ['Sens', 'Bottom-up', 'Top-down'],
        ['Likes / commentaires', 'Oui', 'Non'],
        ['Activable', 'Oui', 'Oui'],
        ['Catégorisation', 'Libre', 'Stricte (7 catégories)'],
        ['Image cover', 'Optionnelle', 'Recommandée'],
        ['Accusé de lecture', 'Non', 'Oui (Phase 3)'],
    ]
)

h2('7.2 Catégories')
bullet("Direction", "communications de la direction")
bullet("RH", "recrutement, formation, paie")
bullet("Vie sociale", "événements, animations")
bullet("Cuisine", "menus, allergènes")
bullet("Protocoles", "nouveaux protocoles soins")
bullet("Sécurité", "exercices feu, alarmes")
bullet("Divers", "fourre-tout")

h2('7.3 Pourquoi pas de likes/commentaires ?')
para("Décision design : une annonce est un acte de communication formel.", bold=True)
bullet("Polluerait le canal (smileys sur « Décès de Mme X »)")
bullet("Donnerait l'illusion d'un dialogue alors que c'est descendant")
bullet("Le mur social existe pour ça")

h2('7.4 Épinglage')
para("Une annonce épinglée reste en haut de la liste. Utile pour les infos urgentes ou stables.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  8. ACCUSÉS DE LECTURE
# ═══════════════════════════════════════════════════════
h1('8. Accusés de lecture (Phase 3)')

h2('8.1 Le problème')
info_box("« Tu as lu le nouveau protocole sur les chutes ? » — « Heu... lequel ? »")
para("Sans accusé de lecture, impossible de prouver qu'une consigne critique a été reçue.")

h2('8.2 Solution')
para("Sur chaque annonce, un toggle « Exiger un accusé de lecture » + un sélecteur de cible (rôle).")
para("Côté lecteur :", bold=True)
bullet("Bandeau jaune en haut de l'annonce : « Cette annonce nécessite votre accusé de lecture »")
bullet("Bouton vert « J'ai lu et compris »")
bullet("Une fois cliqué : bandeau devient vert « Vous avez confirmé la lecture »")
para("Côté responsable :", bold=True)
bullet("Dans l'éditeur d'annonce, dashboard intégré")
bullet("Barre de progression % de lecture")
bullet("Liste verte des collaborateurs ayant lu")
bullet("Liste grise de ceux qui n'ont pas encore lu")
bullet("Permet de relancer ceux qui n'ont pas lu")

h2('8.3 Inspiration')
bullet("SharePoint", "Read receipt")
bullet("Workplace Meta", "Important post")
bullet("Réglementation santé", "traçabilité des consignes")

h2('8.4 Évolutions possibles')
bullet("Rappel automatique par email/notif si pas d'ack après N jours")
bullet("Quizz obligatoire : pour les protocoles très critiques, exiger 3 questions correctes")
bullet("Signature numérique avec date + IP")
bullet("Export PDF de la liste des lecteurs (preuve légale)")
bullet("Limites temporelles : « Cette annonce expire dans 30 jours »")

# ═══════════════════════════════════════════════════════
#  9. RECHERCHE GLOBALE
# ═══════════════════════════════════════════════════════
h1('9. Recherche globale')

h2('9.1 Périmètre')
para("Une seule barre en haut de chaque app cherche dans tout :")
bullet("Wiki (avec filtre permissions)")
bullet("Annonces")
bullet("Résidents")
bullet("Documents")
bullet("Collaborateurs")

h2('9.2 Trois implémentations parallèles')
add_table(
    ['App', 'Endpoint', 'Historique localStorage'],
    [
        ['SpocSpace SPA (employé)', 'global_search', 'spocspace:search-history'],
        ['SpocSpace Admin', 'admin_global_search', 'spocadmin:search-history'],
        ['SpocCare', 'admin_care_global_search', 'spoccare:search-history'],
    ]
)

h2('9.3 Historique avec TTL')
bullet("Les recherches sont conservées 2 jours en localStorage")
bullet("Maximum 30 entrées")
bullet("Affichage : 3 récentes + 4 plus anciennes")
bullet("Bouton X pour supprimer une entrée")
bullet("Icône horloge pour les récentes, loupe pour les plus anciennes")

h2('9.4 Click sur un résultat')
para("Le click ouvre directement l'item. L'id est passé via params.id ou AdminURL.currentId().")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  10. ROADMAP
# ═══════════════════════════════════════════════════════
h1('10. Ce qui manque encore (roadmap)')

h2('10.1 Notifications push')
para("Aujourd'hui : aucune notif quand une nouvelle annonce / page wiki est créée.")
para("À ajouter :", bold=True)
bullet("Notification web push (Service Worker)")
bullet("Email digest hebdomadaire des nouveautés")
bullet("Badge « Nouveau » sur les annonces non lues")

h2('10.2 Templates de pages')
para("Permettre de créer une page à partir d'un template (Procédure, Protocole, FAQ, Note interne...).")

h2('10.3 Mentions @utilisateur')
para("Dans le contenu, pouvoir taper @Marie pour notifier une personne.")

h2('10.4 Liens internes wiki ↔ wiki')
para("Comme Notion : un lien [[Protocole hygiène]] qui crée un lien automatique vers la page.")

h2('10.5 Export PDF de pages')
para("Bouton « Exporter en PDF » pour avoir une copie papier (utile en chambre).")

h2('10.6 Application mobile')
para("PWA installable avec :")
bullet("Mode hors-ligne (cache des pages favorites)")
bullet("Push notifications")
bullet("Scan QR code → ouvre directement une page (QR codes en chambre vers le protocole correspondant)")

h2('10.7 Multi-langue')
para("Une page peut avoir des traductions (FR / EN / PT / ES) pour les équipes internationales.")

h2('10.8 Statistiques personnelles')
para("« Vous avez lu 12 pages cette semaine, 3 de plus que la moyenne » (gamification douce).")

h2('10.9 IA générative complète')
bullet("Génération de brouillons à partir d'un titre")
bullet("Résumé automatique de longues pages")
bullet("Q&A : « Demande à l'IA » qui répond en citant les pages wiki sources")
bullet("Détection automatique de doublons et de contradictions")

h2('10.10 Workflow plus avancé')
bullet("Reviewer dédié (pas n'importe quel responsable)")
bullet("Approbation multi-niveaux (junior → senior → direction)")
bullet("Date de publication différée")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  11. CHOIX TECHNIQUES
# ═══════════════════════════════════════════════════════
h1('11. Choix techniques expliqués')

h2('11.1 Pourquoi pas une SaaS comme Notion / Guru ?')
add_table(
    ['Critère', 'SaaS', 'SpocSpace'],
    [
        ['Coût mensuel', '8-15 €/user', '0 € marginal'],
        ['Données', 'Aux USA', 'En Suisse, sur votre serveur'],
        ['RGPD santé', 'Compliqué (DPA, sous-traitance)', 'Direct'],
        ['Personnalisation', 'Limitée', 'Totale'],
        ['Lien planning/RH/résidents', 'Aucun', 'Natif'],
    ]
)

h2('11.2 Pourquoi MySQL et pas Elasticsearch ?')
bullet("Volume attendu : quelques milliers de pages → MySQL FULLTEXT suffit largement")
bullet("Pas de serveur supplémentaire à maintenir")
bullet("Pas de coût opérationnel")
bullet("À très grosse échelle (>100 000 pages), passage à Meilisearch envisageable")

h2('11.3 Pourquoi vanilla JS et pas React/Vue ?')
bullet("Pas de build step → édition en direct")
bullet("Pas de surface d'attaque npm")
bullet("Démarrage instantané (pas de bundle 500 KB à parser)")
bullet("Maintenance long terme facilitée")

h2('11.4 Pourquoi TipTap et pas CKEditor / TinyMCE ?')
bullet("Open source MIT (gratuit même en commercial)")
bullet("Architecture moderne (ProseMirror)")
bullet("Extensions à la carte")
bullet("Excellente prise en charge des tableaux")
bullet("Sortie HTML clean")
bullet("Communauté active")

# ═══════════════════════════════════════════════════════
#  12. GLOSSAIRE
# ═══════════════════════════════════════════════════════
h1('12. Glossaire')
add_table(
    ['Terme', 'Définition'],
    [
        ['Page', 'Un article wiki avec titre, contenu HTML, métadonnées'],
        ['Catégorie', 'Groupe thématique d\'une page (1 par page)'],
        ['Tag', 'Étiquette transversale (N par page)'],
        ['Expert', 'Utilisateur responsable de la véracité d\'une page'],
        ['Verification', 'Acte de re-valider qu\'une page est encore exacte'],
        ['Carte orpheline', 'Page non lue + non modifiée depuis 90+ jours'],
        ['Knowledge gap', 'Recherche sans résultat → page manquante'],
        ['Brouillon / Review / Publié', 'Les 3 états du workflow de publication'],
        ['Accusé de lecture', 'Confirmation explicite qu\'un user a lu une annonce'],
        ['FULLTEXT', 'Index MySQL pour recherche textuelle rapide'],
    ]
)

# ═══════════════════════════════════════════════════════
#  13. INDEX FICHIERS
# ═══════════════════════════════════════════════════════
h1('13. Index des fichiers source')

h2('Backend')
bullet("admin/api_modules/wiki.php")
bullet("admin/api_modules/annonces.php")
bullet("api_modules/wiki.php")
bullet("admin/api_modules/global_search.php")
bullet("admin/api_modules/care_search.php")
bullet("api_modules/search.php")

h2('Frontend (care)')
bullet("care/pages/wiki.php — liste + lecture")
bullet("care/pages/wiki-edit.php — éditeur")
bullet("care/pages/annonces.php — liste + lecture")
bullet("care/pages/annonce-edit.php — éditeur")

h2('Frontend (SPA employé)')
bullet("pages/wiki.php — vue lecture employé")
bullet("assets/js/modules/wiki.js — module SPA")

h2('Admin')
bullet("admin/pages/wiki.php — wrapper")
bullet("admin/pages/annonces.php — wrapper")
bullet("admin/pages/wiki-analytics.php — dashboard analytics")

h2('Migrations')
bullet("056_wiki.sql — tables initiales")
bullet("057_annonces.sql — annonces + toggle features")
bullet("058_wiki_phase1.sql — tags, favoris, expert, FULLTEXT")
bullet("059_wiki_phase2.sql — permissions, AI suggestions")
bullet("060_wiki_phase3.sql — analytics, search log, ack, workflow")

# ═══════════════════════════════════════════════════════
#  PIED DE PAGE
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run('— Document généré pour SpocSpace —')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY_TEXT

foot2 = doc.add_paragraph()
foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot2.add_run('Phase 1 : Tags / Favoris / Expert / FULLTEXT\nPhase 2 : Permissions par rôle / Suggestions IA\nPhase 3 : Analytics / Knowledge gaps / Accusés de lecture / Workflow brouillon → review → publié')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY_TEXT

# ═══════════════════════════════════════════════════════
#  SAUVEGARDE
# ═══════════════════════════════════════════════════════
output_path = '/home/clients/c81789f8de36e992da19fb6856aa48f6/sites/zkriva.com/spocspace/docs/Wiki_Annonces_Documentation.docx'
doc.save(output_path)
print(f'✓ Document généré : {output_path}')
