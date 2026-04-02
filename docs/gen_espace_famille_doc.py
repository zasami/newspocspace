#!/usr/bin/env python3
"""
Génère le document Word : Espace Famille — Guide complet
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

# ── Fonctions utilitaires ──
def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' — ' + text)
    else:
        p.add_run(text)
    return p

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('ℹ️ ' + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('🏥')
run.font.size = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Espace Famille')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guide complet — Fonctionnement et administration')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SpocSpace — EMS La Terrassière SA')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mars 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Table des matières', level=1)
toc_items = [
    '1. Présentation générale',
    '2. Architecture et sécurité E2EE',
    '3. Guide administration (staff)',
    '   3.1 Accès à l\'espace famille admin',
    '   3.2 Gestion des clés de chiffrement',
    '   3.3 Activités',
    '   3.4 Suivi médical',
    '   3.5 Galerie photos',
    '   3.6 Import par lot',
    '4. Guide famille (correspondants)',
    '   4.1 Connexion',
    '   4.2 Tableau de bord',
    '   4.3 Consulter les activités',
    '   4.4 Consulter le suivi médical',
    '   4.5 Galerie photos et lightbox',
    '5. Réservation restaurant (famille)',
    '   5.1 Accès et fonctionnement',
    '   5.2 Étapes de réservation',
    '   5.3 Menu de la semaine (carousel)',
    '   5.4 Administration des réservations',
    '6. Gestion des résidents',
    '7. Stockage et capacité',
    '8. Sécurité détaillée',
    '9. FAQ / Dépannage',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PRÉSENTATION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Présentation générale', level=1)

doc.add_paragraph(
    'L\'Espace Famille est un module de SpocSpace qui permet aux familles et correspondants '
    'des résidents de suivre la vie quotidienne de leur proche au sein de l\'EMS. '
    'Il offre trois volets principaux :'
)

add_bullet('Suivi des activités et animations (sorties, ateliers, fêtes)', 'Activités')
add_bullet('Consultation des avis médicaux, rapports et ordonnances', 'Suivi médical')
add_bullet('Albums photo organisés par date et par année', 'Galerie photos')

doc.add_paragraph(
    'Tous les fichiers (photos, documents PDF/Word/Excel) sont chiffrés de bout en bout (E2EE). '
    'Le serveur ne voit jamais les fichiers en clair. Seul le correspondant authentifié peut les déchiffrer '
    'dans son navigateur grâce au code d\'accès du résident.'
)

doc.add_heading('Qui fait quoi ?', level=2)

add_table(
    ['Rôle', 'Actions', 'Accès'],
    [
        ['Admin / Direction', 'Créer activités, avis médicaux, albums. Uploader photos et fichiers chiffrés. Gérer les clés E2EE.', 'Panel admin → Espace Famille'],
        ['Responsable / Infirmier', 'Même actions que admin', 'Panel admin → Espace Famille'],
        ['Animateur', 'Créer activités + uploader photos', 'Panel admin → Espace Famille'],
        ['Famille / Correspondant', 'Consulter uniquement (lecture seule)', 'Site web → Espace Famille'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE E2EE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. Architecture et sécurité E2EE', level=1)

doc.add_paragraph(
    'E2EE signifie « End-to-End Encryption » (chiffrement de bout en bout). '
    'Cela signifie que les fichiers sont chiffrés dans le navigateur AVANT d\'être envoyés au serveur, '
    'et déchiffrés dans le navigateur APRÈS avoir été téléchargés. Le serveur ne stocke que des données illisibles.'
)

doc.add_heading('Algorithmes utilisés', level=2)
add_bullet('AES-256-GCM pour le chiffrement des fichiers (standard militaire)', 'Chiffrement')
add_bullet('PBKDF2 avec 100 000 itérations + SHA-256 pour dériver la clé depuis le code d\'accès', 'Dérivation de clé')
add_bullet('IV (vecteur d\'initialisation) unique par fichier, stocké en base', 'Unicité')

doc.add_heading('Flux de chiffrement', level=2)

steps = [
    '1. L\'admin génère une clé AES-256 aléatoire pour un résident (bouton « Générer la clé E2EE »)',
    '2. Cette clé est « enveloppée » (chiffrée) avec le code d\'accès du résident via PBKDF2',
    '3. La clé enveloppée est stockée sur le serveur (illisible sans le code d\'accès)',
    '4. Lors d\'un upload, le fichier est chiffré côté navigateur avec la clé AES du résident',
    '5. Le fichier chiffré (.enc) est envoyé au serveur et stocké sur disque',
    '6. Quand la famille se connecte, elle saisit le code d\'accès → le navigateur dérive la clé',
    '7. Le navigateur télécharge le fichier chiffré et le déchiffre localement',
    '8. L\'image/document s\'affiche — le serveur n\'a jamais vu le contenu en clair',
]
for s in steps:
    doc.add_paragraph(s)

add_note('Si le code d\'accès du résident change, il faut régénérer la clé E2EE dans le panel admin.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. GUIDE ADMIN
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Guide administration (staff)', level=1)

doc.add_heading('3.1 Accès à l\'espace famille admin', level=2)
doc.add_paragraph(
    'Depuis le panel admin, cliquez sur « Espace Famille » dans la sidebar (section Configuration). '
    'Un sélecteur de résident apparaît en haut de la page. Choisissez le résident concerné.'
)

doc.add_heading('3.2 Gestion des clés de chiffrement', level=2)
doc.add_paragraph(
    'Avant de pouvoir uploader des fichiers chiffrés pour un résident, il faut générer sa clé E2EE.'
)
add_bullet('Sélectionnez le résident dans le menu déroulant')
add_bullet('Si le statut affiche « Aucune clé E2EE », cliquez sur « Générer la clé E2EE »')
add_bullet('La clé est générée côté navigateur et enveloppée avec le code d\'accès du résident')
add_bullet('Le statut passe à « Clé E2EE active » (badge vert)')

add_note('La génération de clé nécessite que le résident ait un code d\'accès configuré (auto-généré à la création).')

doc.add_heading('3.3 Activités', level=2)
doc.add_paragraph('Onglet « Activités » dans la page Espace Famille admin.')
add_bullet('Cliquez sur « Nouvelle activité »', 'Créer')
add_bullet('Remplissez le titre (ex: "Sortie Jardin Anglais") et une description', 'Formulaire')
add_bullet('Enregistrez, puis cliquez sur « +Photos » pour ajouter des images', 'Photos')
add_bullet('Les photos sont automatiquement converties en WebP et chiffrées avant l\'envoi', 'Conversion')
add_bullet('Cliquez sur l\'icône crayon pour modifier le titre/description', 'Modifier')
add_bullet('Cliquez sur l\'icône corbeille pour supprimer (photos incluses)', 'Supprimer')

doc.add_heading('3.4 Suivi médical', level=2)
doc.add_paragraph('Onglet « Médical » dans la page Espace Famille admin.')
add_bullet('Cliquez sur « Nouvel avis »')
add_bullet('Remplissez le titre, la date, le type (Avis, Rapport, Ordonnance, Autre)')
add_bullet('Le contenu texte est chiffré côté client avant envoi')
add_bullet('Cliquez sur « +Fichier » pour joindre un PDF, Word, Excel ou image')
add_bullet('Les fichiers sont chiffrés côté navigateur avant l\'envoi au serveur')

add_table(
    ['Type de fichier', 'Extensions supportées', 'Icône'],
    [
        ['PDF', '.pdf', '📄 Rouge'],
        ['Word', '.doc, .docx', '📄 Bleu'],
        ['Excel', '.xls, .xlsx', '📄 Vert'],
        ['Image', '.jpg, .jpeg, .png', '📄 Violet'],
    ]
)

doc.add_heading('3.5 Galerie photos', level=2)
doc.add_paragraph('Onglet « Galerie » dans la page Espace Famille admin.')
add_bullet('Cliquez sur « Nouvel album » — donnez un titre, une date et une année', 'Créer un album')
add_bullet('Cliquez sur « +Photos » pour ajouter des images à l\'album', 'Ajouter des photos')
add_bullet('Les photos sont converties en WebP + chiffrées automatiquement', 'Conversion')
add_bullet('La première photo devient automatiquement la couverture de l\'album', 'Couverture')

doc.add_heading('3.6 Import par lot', level=2)
doc.add_paragraph(
    'Le bouton « Import par lot » dans l\'onglet Galerie permet d\'importer plusieurs photos en une fois. '
    'Le système :'
)
add_bullet('Lit le nom de fichier ou la date de modification de chaque photo')
add_bullet('Regroupe les photos par mois (ex: "Mars 2026", "Avril 2026")')
add_bullet('Crée automatiquement un album par mois')
add_bullet('Convertit chaque photo en WebP, la chiffre, et l\'uploade')
add_bullet('L\'admin peut ensuite renommer les albums depuis la liste')

add_note('Les noms de fichiers contenant une date (ex: IMG_20260315_001.jpg) sont automatiquement détectés.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. GUIDE FAMILLE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Guide famille (correspondants)', level=1)

doc.add_heading('4.1 Connexion', level=2)
doc.add_paragraph(
    'Accédez à l\'Espace Famille depuis le site web de l\'EMS (bouton « Espace Famille » '
    'dans le widget latéral) ou directement via l\'URL :'
)
p = doc.add_paragraph()
run = p.add_run('https://zkriva.com/spocspace/website/famille.php')
run.bold = True

doc.add_paragraph('Saisissez :')
add_bullet('L\'email du correspondant (fourni par l\'administration)')
add_bullet('Le code d\'accès = date de naissance du résident au format JJMMAAAA (ex: 12031935)')

add_note('Après 5 tentatives échouées, l\'accès est bloqué pendant 15 minutes (protection anti-brute-force).')

doc.add_heading('4.2 Tableau de bord', level=2)
doc.add_paragraph(
    'Après connexion, la sidebar à gauche affiche la photo, le nom et la chambre du résident. '
    'Le tableau de bord (Accueil) présente :'
)
add_bullet('Les statistiques : nombre d\'activités, avis médicaux, albums et photos')
add_bullet('Les dernières activités avec titre et date')
add_bullet('Les derniers avis médicaux avec type et date')

doc.add_heading('4.3 Consulter les activités', level=2)
doc.add_paragraph('Cliquez sur « Activités » dans la sidebar.')
add_bullet('Liste des activités avec date, titre et description')
add_bullet('Badge indiquant le nombre de photos')
add_bullet('Cliquez sur une activité pour voir le détail et les photos')
add_bullet('Les photos sont déchiffrées automatiquement dans votre navigateur')
add_bullet('Cliquez sur une photo pour l\'agrandir en plein écran (lightbox)')

doc.add_heading('4.4 Consulter le suivi médical', level=2)
doc.add_paragraph('Cliquez sur « Suivi médical » dans la sidebar.')
add_bullet('Timeline des avis médicaux triés par date')
add_bullet('Chaque avis affiche son type (Avis, Rapport, Ordonnance) avec un code couleur')
add_bullet('Le contenu texte chiffré est déchiffré et affiché automatiquement')
add_bullet('Cliquez sur un fichier joint pour le visualiser ou le télécharger')
add_bullet('Les PDF s\'ouvrent dans la lightbox — les autres fichiers se téléchargent')

add_table(
    ['Type', 'Couleur de la bordure', 'Icône'],
    [
        ['Avis', 'Vert', '📋'],
        ['Rapport', 'Bleu', '📊'],
        ['Ordonnance', 'Orange', '💊'],
        ['Autre', 'Marron', '📄'],
    ]
)

doc.add_heading('4.5 Galerie photos et lightbox', level=2)
doc.add_paragraph('Cliquez sur « Galerie » dans la sidebar.')
add_bullet('Les albums sont organisés par année avec des séparateurs')
add_bullet('Chaque album affiche une couverture (première photo), le titre, la date et le nombre de photos')
add_bullet('Cliquez sur un album pour voir toutes les photos en grille')
add_bullet('Cliquez sur une photo pour ouvrir la lightbox plein écran')
add_bullet('Navigation : flèches gauche/droite ou touches clavier ← →')
add_bullet('Fermer : bouton × ou touche Escape')
add_bullet('Télécharger : bouton de téléchargement en haut à gauche')

add_note('Les photos sont déchiffrées progressivement (lazy loading). Si beaucoup de photos, '
         'un spinner s\'affiche pendant le déchiffrement.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. RÉSERVATION RESTAURANT
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. Réservation restaurant (famille)', level=1)

doc.add_paragraph(
    'Le site web public de l\'EMS permet aux familles et correspondants de réserver un repas '
    'pour accompagner leur proche résident. La réservation se fait directement depuis la page '
    'du menu de la semaine.'
)

doc.add_heading('5.1 Accès et fonctionnement', level=2)

doc.add_paragraph(
    'Le module de réservation est accessible depuis le site web principal de l\'EMS :'
)
p = doc.add_paragraph()
run = p.add_run('https://zkriva.com/spocspace/website/')
run.bold = True

doc.add_paragraph(
    'La section « Menu de la semaine » affiche un carousel de cartes jour par jour, '
    'avec les menus midi et soir pour chaque journée. Chaque carte dispose d\'un bouton '
    '« Réserver un repas » (désactivé pour les dates passées).'
)

doc.add_heading('5.2 Étapes de réservation', level=2)

steps_resa = [
    '1. Cliquez sur « Réserver un repas » sur la carte du jour souhaité',
    '2. Un modal s\'ouvre — saisissez l\'email du correspondant et le code d\'accès (date de naissance JJMMAAAA)',
    '3. Le système identifie le résident associé et affiche ses informations',
    '4. Choisissez le repas : Midi ou Soir',
    '5. Indiquez le nombre de personnes (1 à 5)',
    '6. Ajoutez des remarques si nécessaire (allergies, régime...)',
    '7. Cliquez sur « Confirmer la réservation »',
    '8. Un ticket de confirmation s\'affiche avec les détails',
]
for s in steps_resa:
    doc.add_paragraph(s)

add_note('Une fois connecté, le correspondant reste identifié pour les réservations suivantes '
         '(pas besoin de ressaisir email/code). La session est conservée dans le navigateur.')

doc.add_heading('Tarifs', level=3)

add_table(
    ['Repas', 'Prix'],
    [
        ['Midi', 'CHF 14.50'],
        ['Soir', 'CHF 11.00'],
    ]
)

doc.add_heading('Règles de réservation', level=3)
add_bullet('Impossible de réserver pour une date passée')
add_bullet('Une seule réservation par résident, par repas, par date')
add_bullet('Le correspondant ne voit que le résident qui lui est associé')

doc.add_heading('5.3 Menu de la semaine (carousel)', level=2)

doc.add_paragraph(
    'Le carousel affiche les menus sur 4 semaines (28 jours) dans un défilement continu :'
)
add_bullet('3 cartes visibles simultanément (jours consécutifs)')
add_bullet('Flèches gauche/droite : avancer ou reculer d\'1 jour')
add_bullet('Passage automatique d\'une semaine à l\'autre en continuité')
add_bullet('Boutons « semaine précédente / suivante » pour sauter de 7 jours')
add_bullet('Le jour actuel est mis en surbrillance avec un badge « Aujourd\'hui »')

doc.add_paragraph('Chaque carte affiche :')
add_bullet('Le jour et la date')
add_bullet('Menu du midi : entrée, plat, accompagnement, salade, dessert')
add_bullet('Menu du soir : même structure')
add_bullet('Prix par repas')
add_bullet('Remarques éventuelles du chef')

doc.add_heading('5.4 Administration des réservations', level=2)

doc.add_paragraph(
    'Les réservations famille sont gérées côté admin dans la page « Réservations repas » :'
)
add_bullet('Onglet « Collaborateurs » : réservations des employés de l\'EMS', 'Tab 1')
add_bullet('Onglet « Famille / Visiteurs » : réservations des correspondants de résidents', 'Tab 2')

doc.add_paragraph('Pour chaque réservation, l\'admin peut voir :')
add_bullet('Le nom du résident et du visiteur')
add_bullet('La chambre du résident')
add_bullet('Le choix (menu ou salade)')
add_bullet('Le nombre de personnes')
add_bullet('Les remarques (allergies, régime)')
add_bullet('Le créateur de la réservation')

doc.add_paragraph(
    'Un bouton « Imprimer » permet de générer une version imprimable des réservations du jour, '
    'utile pour la cuisine.'
)

doc.add_heading('Administration des menus', level=3)
doc.add_paragraph(
    'Les menus sont gérés dans la page « Menus » du panel admin (section Cuisine). '
    'Les animateurs cuisine créent les menus midi et soir pour chaque jour. '
    'Les menus sont automatiquement affichés sur le site web public.'
)

add_note('Les menus du site web sont mis en cache côté serveur (SSR). '
         'Ils sont chargés au rendu de la page, pas en AJAX, '
         'pour un affichage instantané sans spinner.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. GESTION DES RÉSIDENTS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Gestion des résidents', level=1)

doc.add_paragraph(
    'La page « Résidents » dans le panel admin permet de gérer tous les résidents de l\'EMS.'
)

doc.add_heading('Informations d\'un résident', level=2)
add_table(
    ['Champ', 'Description'],
    [
        ['Nom / Prénom', 'Identité du résident'],
        ['Date de naissance', 'Utilisée pour le code d\'accès famille (JJMMAAAA)'],
        ['Chambre / Étage', 'Localisation dans l\'EMS'],
        ['Photo', 'Photo de profil (affichée dans la liste et l\'espace famille)'],
        ['Correspondant', 'Nom, prénom, email et téléphone du contact famille'],
        ['Code d\'accès', 'Auto-généré (nom + chambre). Utilisé aussi comme mot de passe famille'],
        ['VIP', 'Résident avec menu spécial'],
        ['Actif', 'Permet de désactiver sans supprimer'],
    ]
)

doc.add_heading('Photo du résident', level=2)
add_bullet('Cliquez sur une ligne dans la liste des résidents pour ouvrir le modal d\'édition')
add_bullet('Cliquez sur la zone photo ronde en haut du modal')
add_bullet('Sélectionnez une image — elle est convertie en WebP automatiquement')
add_bullet('La photo s\'affiche dans la liste admin ET dans la sidebar de l\'espace famille')
add_bullet('Pour supprimer : survolez la photo → bouton × rouge')

doc.add_heading('Recherche', level=2)
doc.add_paragraph(
    'Utilisez la barre de recherche globale (topbar) pour filtrer les résidents par nom, prénom ou chambre. '
    'La recherche est instantanée (live search).'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. STOCKAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Stockage et capacité', level=1)

doc.add_paragraph(
    'Les fichiers chiffrés sont stockés dans le dossier uploads/famille/ sur le serveur, '
    'organisés par résident. L\'accès direct est bloqué par .htaccess.'
)

doc.add_heading('Estimation de capacité (1 To d\'espace disque)', level=2)

add_table(
    ['Paramètre', 'Valeur'],
    [
        ['Taille moyenne photo WebP', '150-300 Ko'],
        ['Photos/jour (94 résidents × 10 photos)', '940 photos/jour'],
        ['Espace/mois', '~7 Go'],
        ['Espace/an', '~86 Go'],
        ['Durée avec 1 To', '~11 ans'],
        ['En pratique (2-3 photos/résident/jour)', '30+ ans'],
    ]
)

doc.add_heading('Organisation des fichiers', level=2)
doc.add_paragraph('uploads/famille/{resident_id}/')
add_bullet('activites/ — Photos des activités')
add_bullet('medical/ — Fichiers médicaux (PDF, Word, Excel, images)')
add_bullet('galerie/ — Photos des albums')

add_note('Tous les fichiers sont stockés sous un nom UUID avec extension .enc. '
         'Le nom original est conservé en base de données.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. SÉCURITÉ
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('8. Sécurité détaillée', level=1)

add_table(
    ['Mesure', 'Description'],
    [
        ['E2EE (AES-256-GCM)', 'Tous les fichiers chiffrés côté client, serveur ne voit jamais le contenu'],
        ['PBKDF2 (100K itérations)', 'Dérivation de clé résistante au brute-force'],
        ['Sessions tokenisées', 'Token unique 96 caractères, expire en 24h'],
        ['Rate limiting', '5 tentatives max par IP, blocage 15 minutes'],
        ['.htaccess Deny', 'Accès direct aux fichiers uploadés interdit'],
        ['CSP (Content Security Policy)', 'blob: autorisé pour les images déchiffrées uniquement'],
        ['Cache-Control: no-store', 'Fichiers déchiffrés non mis en cache par le navigateur'],
        ['UUID pour noms de fichiers', 'Noms originaux masqués sur le disque'],
        ['Vérification résident_id', 'Chaque requête vérifie que le correspondant a accès au bon résident'],
        ['CSRF protection', 'Token CSRF sur toutes les actions admin'],
        ['Conversion WebP', 'Réduction de taille + suppression des métadonnées EXIF (privacy)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. FAQ
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('9. FAQ / Dépannage', level=1)

faqs = [
    ('La famille ne peut pas se connecter',
     'Vérifiez que l\'email du correspondant est correct dans la fiche résident. '
     'Le code d\'accès est la date de naissance au format JJMMAAAA (ex: 12031935 pour le 12 mars 1935). '
     'Si 5 tentatives échouées, attendez 15 minutes.'),
    ('Les photos ne s\'affichent pas (icône cadenas)',
     'La clé E2EE n\'a pas été générée pour ce résident. '
     'Allez dans Admin → Espace Famille → sélectionnez le résident → cliquez « Générer la clé E2EE ».'),
    ('Le message « Clé E2EE non disponible » apparaît',
     'La famille s\'est reconnectée sans saisir le code d\'accès (session restaurée par token). '
     'Le code d\'accès est nécessaire pour déchiffrer la clé. '
     'Solution : se déconnecter et se reconnecter en saisissant le code.'),
    ('J\'ai changé le code d\'accès d\'un résident',
     'Il faut régénérer la clé E2EE dans Admin → Espace Famille. '
     'ATTENTION : les fichiers chiffrés avec l\'ancienne clé ne seront plus lisibles. '
     'Il faudra les re-uploader.'),
    ('Combien de photos puis-je uploader à la fois ?',
     'Pas de limite technique. L\'import par lot gère des centaines de photos. '
     'Chaque photo est convertie en WebP (~250 Ko) et chiffrée avant envoi. '
     'Comptez ~1-2 secondes par photo selon la connexion.'),
    ('Les fichiers PDF médicaux sont-ils sécurisés ?',
     'Oui, ils sont chiffrés avec le même algorithme AES-256-GCM que les photos. '
     'Le serveur ne stocke que des fichiers .enc illisibles. '
     'Seul le correspondant avec le bon code d\'accès peut les ouvrir.'),
    ('Puis-je accéder à l\'Espace Famille depuis un téléphone ?',
     'Oui, l\'interface est entièrement responsive. '
     'Sur mobile, la sidebar se transforme en menu hamburger.'),
    ('Comment supprimer un résident et ses données ?',
     'Désactivez le résident dans Admin → Résidents (bouton toggle). '
     'Les données famille restent en base mais ne sont plus accessibles. '
     'La suppression définitive nécessite une intervention technique.'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run('Q : ' + q)
    run.bold = True
    p = doc.add_paragraph('R : ' + a)
    p.paragraph_format.space_after = Pt(12)

# ── Footer ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Document généré automatiquement par SpocSpace — Mars 2026 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── Save ──
output = '/home/clients/c81789f8de36e992da19fb6856aa48f6/sites/zkriva.com/spocspace/docs/Espace_Famille_Guide_Complet.docx'
doc.save(output)
print(f'Document saved: {output}')
