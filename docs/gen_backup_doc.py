#!/usr/bin/env python3
"""
Genere le document Word : Sauvegarde & Restauration — Documentation complete
Usage : python3 gen_backup_doc.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# -- Marges --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

CARE_GREEN = RGBColor(0x2D, 0x4A, 0x43)
ACCENT_ORANGE = RGBColor(0xC5, 0x4A, 0x3A)
DANGER_RED = RGBColor(0xCC, 0x33, 0x33)
GREY_TEXT = RGBColor(0x6B, 0x6B, 0x6B)
BLUE_INFO = RGBColor(0x2B, 0x6C, 0xB0)

for level, size in [(1, 22), (2, 16), (3, 13), (4, 11)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = CARE_GREEN
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.bold = True

# -- Helpers --
def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

def para(text, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if italic: r.italic = True
    if bold: r.bold = True
    if color: r.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(' — ' + text)
    else:
        p.add_run(text)
    return p

def info_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY_TEXT

def danger_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = DANGER_RED

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return table

# ============================================================
# PAGE DE TITRE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('SpocSpace', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.size = Pt(28)

subtitle = doc.add_heading('Sauvegarde & Restauration', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in subtitle.runs:
    r.font.size = Pt(20)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documentation technique et guide utilisateur')
r.font.size = Pt(13)
r.font.color.rgb = GREY_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 1.0 — Avril 2026')
r.font.size = Pt(11)
r.font.color.rgb = GREY_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EMS Terrassiere — Geneve')
r.font.size = Pt(11)
r.font.color.rgb = GREY_TEXT

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES
# ============================================================
h1('Table des matieres')
toc_items = [
    '1. Introduction',
    '2. Architecture du systeme de sauvegarde',
    '   2.1 Vue d\'ensemble',
    '   2.2 Format de stockage (ZIP)',
    '   2.3 Structure des fichiers',
    '   2.4 Base de donnees — Table backups',
    '3. Sauvegardes par utilisateur (Per-user)',
    '   3.1 Principe et perimetre',
    '   3.2 Declenchement',
    '   3.3 Frequence et retention',
    '   3.4 Guide utilisateur — Creer une sauvegarde',
    '4. Sauvegardes globales',
    '   4.1 Principe et perimetre',
    '   4.2 Sauvegarde automatique (Cron)',
    '   4.3 Sauvegarde manuelle',
    '   4.4 Frequence et retention',
    '   4.5 Code d\'acces special',
    '5. Restauration',
    '   5.1 Restauration per-user — Comparaison',
    '   5.2 Restauration per-user — Ecrasement',
    '   5.3 Restauration globale',
    '   5.4 Flux UX detaille',
    '6. Securite',
    '7. API — Reference technique',
    '8. Estimation espace disque',
    '9. FAQ',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
h1('1. Introduction')
para('Le module Sauvegarde & Restauration de SpocSpace permet aux administrateurs de '
     'proteger les donnees de l\'application contre la perte accidentelle, les erreurs '
     'humaines et les incidents techniques.')
para('Il offre deux niveaux de protection :')
bullet('Sauvegardes par utilisateur (per-user)', 'Niveau 1')
bullet('Sauvegardes globales du systeme complet', 'Niveau 2')
para('Ce document couvre l\'architecture technique, le guide d\'utilisation pour les '
     'administrateurs, les procedures de restauration, et les aspects de securite.')

# ============================================================
# 2. ARCHITECTURE
# ============================================================
h1('2. Architecture du systeme de sauvegarde')

h2('2.1 Vue d\'ensemble')
para('Le systeme repose sur des archives ZIP generees par PHP (extension ZipArchive native). '
     'Chaque sauvegarde contient un dump SQL des donnees pertinentes, les fichiers physiques '
     'associes (documents uploades), et des metadonnees de verification.')

h2('2.2 Format de stockage — ZIP')
para('Le format ZIP a ete choisi pour les raisons suivantes :')
bullet('Support natif PHP via ZipArchive — aucune extension externe requise')
bullet('Peut combiner dump SQL + fichiers physiques dans une seule archive')
bullet('Bon ratio de compression sur du texte SQL/JSON (80-90% de reduction)')
bullet('Format standard, facilement inspectable et verifiable')
bullet('Compatible avec tous les systemes d\'exploitation')

info_box('Alternative evaluee : tar.gz offre une meilleure compression mais necessite '
         'des extensions externes en PHP. gzip seul ne gere qu\'un fichier unique.')

h2('2.3 Structure des fichiers')
para('Les sauvegardes sont stockees dans un repertoire dedie, protege contre l\'acces web :')

para('data/backups/', bold=True)
p = doc.add_paragraph()
r = p.add_run(
    'data/backups/\n'
    '  .htaccess              (Deny from all)\n'
    '  users/\n'
    '    {user_id}/\n'
    '      2026-04-16_143022.zip\n'
    '      2026-04-14_090000.zip\n'
    '  global/\n'
    '    daily_2026-04-16.zip\n'
    '    daily_2026-04-15.zip\n'
    '    weekly_2026-W15.zip'
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

h3('Contenu d\'une archive per-user')
add_table(
    ['Fichier', 'Description'],
    [
        ['manifest.json', 'Metadonnees : date, version app, tables incluses, checksums'],
        ['documents.sql', 'INSERT statements des documents de l\'utilisateur'],
        ['messages.sql', 'INSERT statements des messages (inbox + envoyes + supprimes)'],
        ['emails.sql', 'INSERT statements des emails'],
        ['files/', 'Copies des fichiers physiques uploades (PDF, images, etc.)'],
        ['checksum.sha256', 'Hash SHA-256 de chaque fichier pour verification d\'integrite'],
    ]
)

doc.add_paragraph()
h3('Contenu d\'une archive globale')
add_table(
    ['Fichier', 'Description'],
    [
        ['manifest.json', 'Metadonnees completes du systeme'],
        ['full_dump.sql', 'Dump complet de TOUTES les tables de la base de donnees'],
        ['files/', 'Tous les fichiers uploades du systeme'],
        ['checksum.sha256', 'Hash SHA-256 pour verification d\'integrite'],
    ]
)

h2('2.4 Base de donnees — Table backups')
para('Chaque sauvegarde est enregistree dans la table backups pour le suivi et la gestion :')
add_table(
    ['Colonne', 'Type', 'Description'],
    [
        ['id', 'CHAR(36) PK', 'UUID unique'],
        ['user_id', 'CHAR(36) NULL', 'NULL = sauvegarde globale'],
        ['type', 'ENUM(user, global)', 'Type de sauvegarde'],
        ['filename', 'VARCHAR(255)', 'Nom du fichier ZIP'],
        ['file_size', 'INT UNSIGNED', 'Taille en octets'],
        ['tables_included', 'JSON', 'Liste des tables incluses'],
        ['row_counts', 'JSON', 'Nombre de lignes par table'],
        ['checksum_sha256', 'CHAR(64)', 'Hash d\'integrite de l\'archive'],
        ['created_at', 'DATETIME', 'Date de creation'],
        ['created_by', 'CHAR(36)', 'Admin ayant declenche la sauvegarde'],
    ]
)

doc.add_page_break()

# ============================================================
# 3. SAUVEGARDES PER-USER
# ============================================================
h1('3. Sauvegardes par utilisateur (Per-user)')

h2('3.1 Principe et perimetre')
para('Chaque administrateur peut creer une sauvegarde de ses propres donnees. '
     'Le perimetre inclut :')
bullet('Documents personnels et partages')
bullet('Messages envoyes, recus et supprimes')
bullet('Emails et correspondance')
bullet('Fichiers uploades associes')

info_box('Les sauvegardes per-user ne contiennent PAS les donnees systeme (plannings, '
         'configurations, comptes utilisateurs). Pour cela, voir les sauvegardes globales.')

h2('3.2 Declenchement')
para('Les sauvegardes per-user sont declenchees manuellement par l\'administrateur '
     'via le bouton "Creer une sauvegarde" dans la page Sauvegardes du panneau admin.')
para('Cas d\'usage recommandes :', bold=True)
bullet('Avant une modification importante des donnees')
bullet('En fin de mois pour archivage')
bullet('Avant une migration ou mise a jour du systeme')
bullet('Apres une importation de donnees en masse')

h2('3.3 Frequence et retention')
add_table(
    ['Parametre', 'Valeur', 'Justification'],
    [
        ['Declenchement', 'Manuel', 'L\'admin decide quand sauvegarder'],
        ['Maximum par user', '5 sauvegardes', 'Rotation auto : la plus ancienne est supprimee'],
        ['Taille estimee', '2-10 Mo par archive', 'Depend du volume de documents'],
        ['Espace max par user', '~50 Mo', '5 archives x 10 Mo max'],
    ]
)

h2('3.4 Guide utilisateur — Creer une sauvegarde')
para('Etape 1 :', bold=True)
para('Naviguer vers le panneau d\'administration > menu "Sauvegardes".')
para('Etape 2 :', bold=True)
para('Dans l\'onglet "Mes sauvegardes", cliquer sur le bouton "+ Creer une sauvegarde maintenant".')
para('Etape 3 :', bold=True)
para('Le systeme genere l\'archive ZIP. Une barre de progression s\'affiche pendant le traitement. '
     'Un message de confirmation apparait une fois la sauvegarde terminee.')
para('Etape 4 :', bold=True)
para('La sauvegarde apparait dans la liste avec la date, l\'heure et le nombre d\'elements '
     'sauvegardes (documents, messages, emails).')

info_box('Si le nombre maximum de 5 sauvegardes est atteint, la plus ancienne sera '
         'automatiquement supprimee pour faire place a la nouvelle.')

doc.add_page_break()

# ============================================================
# 4. SAUVEGARDES GLOBALES
# ============================================================
h1('4. Sauvegardes globales')

h2('4.1 Principe et perimetre')
para('Les sauvegardes globales capturent l\'integralite du systeme SpocSpace :')
bullet('Toutes les tables de la base de donnees (users, plannings, absences, etc.)')
bullet('Tous les fichiers uploades par tous les utilisateurs')
bullet('La configuration du systeme (ems_config)')
para('Elles permettent une restauration complete du systeme a un point donne dans le temps.')

h2('4.2 Sauvegarde automatique (Cron)')
para('Un script PHP execute automatiquement une sauvegarde globale chaque nuit a 3h00 :')
p = doc.add_paragraph()
r = p.add_run('0 3 * * * php /path/to/spocspace/scripts/backup_daily.php')
r.font.name = 'Consolas'
r.font.size = Pt(9)

para('Le script gere egalement la rotation automatique (suppression des archives obsoletes).')

info_box('Si l\'hebergeur ne supporte pas les taches cron, un mecanisme de pseudo-cron '
         'est disponible : a chaque visite admin, PHP verifie si le dernier backup global '
         'date de plus de 24h et le lance en arriere-plan.')

h2('4.3 Sauvegarde manuelle')
para('Un administrateur peut declencher une sauvegarde globale immediate depuis l\'onglet '
     '"Global" de la page Sauvegardes. Ceci est recommande avant :')
bullet('Une migration de base de donnees')
bullet('Une mise a jour majeure de l\'application')
bullet('Un changement de configuration important')

h2('4.4 Frequence et retention')
add_table(
    ['Type', 'Frequence', 'Retention', 'Espace estime'],
    [
        ['Quotidien', '1x/jour (3h00)', '14 jours glissants', '50-200 Mo x 14 = 0.7-2.8 Go'],
        ['Hebdomadaire', '1x/semaine (dimanche)', '8 semaines', '50-200 Mo x 8 = 0.4-1.6 Go'],
        ['Manuel', 'A la demande', 'Compte dans les quotidiens', 'Inclus dans la rotation'],
    ]
)
doc.add_paragraph()
para('Espace total estime : 3 a 5 Go maximum pour un EMS de taille moyenne.', bold=True)

h2('4.5 Code d\'acces special')
para('La restauration globale est protegee par un code d\'acces special, distinct du mot '
     'de passe administrateur. Ce code est :')
bullet('Defini lors de la configuration initiale du systeme')
bullet('Stocke sous forme hashee dans la table ems_config')
bullet('Rate-limite a 3 tentatives par heure')
bullet('Requis pour CHAQUE restauration globale (pas de "se souvenir")')

danger_box('IMPORTANT : Ce code doit etre connu uniquement de la direction et conserve '
           'en lieu sur (coffre-fort numerique, enveloppe scellee). Sa perte necessite '
           'une intervention technique pour le reinitialiser.')

doc.add_page_break()

# ============================================================
# 5. RESTAURATION
# ============================================================
h1('5. Restauration')

h2('5.1 Restauration per-user — Mode Comparaison')
para('Le mode Comparaison permet de voir les differences entre la sauvegarde et l\'etat actuel '
     'avant de decider quoi restaurer.')
para('Fonctionnement :', bold=True)
bullet('Le systeme extrait l\'archive ZIP et compare chaque table ligne par ligne')
bullet('Un rapport de differences est affiche :')

add_table(
    ['Type', 'Couleur', 'Description'],
    [
        ['+ Ajoutes', 'Vert', 'Elements presents dans la sauvegarde mais absents actuellement'],
        ['- Supprimes', 'Rouge', 'Elements presents actuellement mais absents de la sauvegarde'],
        ['~ Modifies', 'Orange', 'Elements presents des deux cotes mais avec des differences'],
    ]
)

doc.add_paragraph()
para('L\'administrateur peut alors choisir de restaurer uniquement les elements manquants '
     '(les "+") sans toucher aux donnees actuelles. C\'est le mode le plus sur.')

info_box('Ce mode est recommande dans la majorite des cas. Il permet de recuperer des '
         'documents ou messages supprimes sans risquer de perdre des donnees recentes.')

h2('5.2 Restauration per-user — Mode Ecrasement')
para('Le mode Ecrasement remplace TOUTES les donnees actuelles par celles de la sauvegarde.')

danger_box('DANGER : Toutes les donnees actuelles non presentes dans la sauvegarde seront '
           'DEFINITIVEMENT PERDUES. Les modifications effectuees apres la date de la '
           'sauvegarde seront ecrasees sans possibilite de recuperation.')

para('Procedure de securite :', bold=True)
bullet('Un avertissement rouge s\'affiche avec le detail des consequences')
bullet('L\'administrateur doit taper le mot "RESTAURER" dans un champ de confirmation')
bullet('Le bouton de confirmation n\'est actif que si le mot est correctement saisi')
bullet('Une sauvegarde automatique de l\'etat actuel est creee avant l\'ecrasement')

h2('5.3 Restauration globale')
para('La restauration globale remet le systeme ENTIER dans l\'etat de la sauvegarde selectionnee.')

danger_box('DANGER CRITIQUE : Une restauration globale ecrase TOUTES les donnees de TOUS '
           'les utilisateurs. Toutes les modifications effectuees depuis la date de la '
           'sauvegarde seront perdues pour TOUS les utilisateurs du systeme.')

para('Double securite requise :', bold=True)
bullet('Saisie du code d\'acces special (voir section 4.5)', 'Etape 1')
bullet('Saisie du mot "RESTAURER" dans le champ de confirmation', 'Etape 2')
para('Une sauvegarde automatique de l\'etat actuel est TOUJOURS creee avant toute '
     'restauration globale, permettant de revenir en arriere en cas d\'erreur.')

h2('5.4 Flux UX detaille')

h3('Page Sauvegardes — Onglet "Mes sauvegardes"')
para('La page presente la liste des sauvegardes existantes, ordonnees par date decroissante. '
     'Chaque entree affiche :')
bullet('Date et heure de creation')
bullet('Nombre d\'elements par categorie (documents, messages, emails)')
bullet('Taille de l\'archive')
bullet('Trois boutons d\'action : Comparer, Restaurer, Supprimer')

h3('Flux "Comparer"')
bullet('Clic sur "Comparer"', 'Etape 1')
bullet('Le systeme analyse les differences (peut prendre quelques secondes)', 'Etape 2')
bullet('Affichage du rapport : +N ajoutes, -N supprimes, ~N modifies', 'Etape 3')
bullet('Bouton "Restaurer seulement les differences" ou "Annuler"', 'Etape 4')

h3('Flux "Restaurer" (ecrasement)')
bullet('Clic sur "Restaurer"', 'Etape 1')
bullet('Modal d\'avertissement DANGER avec texte en rouge', 'Etape 2')
bullet('Champ de saisie : taper "RESTAURER" pour confirmer', 'Etape 3')
bullet('Sauvegarde auto de l\'etat actuel + restauration', 'Etape 4')
bullet('Message de confirmation avec resume des operations', 'Etape 5')

h3('Onglet "Global"')
bullet('Saisie du code d\'acces special pour acceder a l\'onglet', 'Etape 1')
bullet('Liste des sauvegardes globales (quotidiennes + hebdomadaires)', 'Etape 2')
bullet('Memes actions (Comparer / Restaurer) avec double confirmation', 'Etape 3')

doc.add_page_break()

# ============================================================
# 6. SECURITE
# ============================================================
h1('6. Securite')

h2('Protection du stockage')
bullet('Repertoire data/backups/ protege par .htaccess (Deny from all)', 'Acces web')
bullet('Idealement place hors du webroot du serveur', 'Emplacement')
bullet('Seul le processus PHP peut lire/ecrire les fichiers de sauvegarde', 'Permissions')

h2('Integrite des donnees')
bullet('Chaque archive contient un fichier checksum.sha256', 'Checksums')
bullet('Le hash est verifie automatiquement avant toute restauration', 'Verification')
bullet('Si le hash ne correspond pas, la restauration est refusee', 'Rejet')

h2('Controle d\'acces')
add_table(
    ['Action', 'Niveau requis', 'Protection supplementaire'],
    [
        ['Creer backup per-user', 'Admin', 'Aucune'],
        ['Lister ses backups', 'Admin', 'Aucune'],
        ['Comparer un backup', 'Admin', 'Aucune'],
        ['Restaurer per-user (merge)', 'Admin', 'Confirmation visuelle'],
        ['Restaurer per-user (ecrasement)', 'Admin', 'Saisir "RESTAURER"'],
        ['Creer backup global', 'Admin', 'Aucune'],
        ['Restaurer global', 'Admin', 'Code special + saisir "RESTAURER"'],
    ]
)

h2('Protection contre les abus')
bullet('Maximum 5 sauvegardes per-user (evite la saturation disque)')
bullet('Rate-limiting : 3 tentatives de code special par heure')
bullet('Sauvegarde auto avant ecrasement (filet de securite)')
bullet('Logs de toutes les operations de restauration')

doc.add_page_break()

# ============================================================
# 7. API REFERENCE
# ============================================================
h1('7. API — Reference technique')

para('Toutes les actions sont accessibles via POST sur admin/api.php. '
     'Le module est enregistre dans admin/api_modules/_routes.php.')

h2('Actions disponibles')
add_table(
    ['Action', 'Description', 'Parametres'],
    [
        ['admin_create_backup', 'Creer une sauvegarde per-user', 'Aucun (user courant)'],
        ['admin_list_backups', 'Lister les sauvegardes', 'type: user|global'],
        ['admin_compare_backup', 'Comparer avec etat actuel', 'backup_id'],
        ['admin_restore_backup', 'Restaurer (merge ou ecrasement)', 'backup_id, mode: merge|overwrite'],
        ['admin_delete_backup', 'Supprimer une sauvegarde', 'backup_id'],
        ['admin_create_global_backup', 'Sauvegarde globale manuelle', 'Aucun'],
        ['admin_restore_global_backup', 'Restauration globale', 'backup_id, access_code'],
    ]
)

h2('Reponses')
para('Toutes les reponses suivent le format standard SpocSpace :')
p = doc.add_paragraph()
r = p.add_run(
    '// Succes\n'
    '{ "ok": true, "data": { ... } }\n\n'
    '// Erreur\n'
    '{ "ok": false, "message": "Description de l\'erreur" }'
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

h2('Exemple — Creer une sauvegarde')
p = doc.add_paragraph()
r = p.add_run(
    'POST /spocspace/admin/api.php\n'
    'Content-Type: application/json\n'
    'X-CSRF-Token: {token}\n\n'
    '{ "action": "admin_create_backup" }\n\n'
    '// Reponse\n'
    '{\n'
    '  "ok": true,\n'
    '  "data": {\n'
    '    "id": "a1b2c3d4-...",\n'
    '    "filename": "backup_user_xxx_2026-04-16_143022.zip",\n'
    '    "file_size": 4521984,\n'
    '    "row_counts": { "documents": 42, "messages": 156, "emails": 23 }\n'
    '  }\n'
    '}'
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# 8. ESTIMATION ESPACE
# ============================================================
h1('8. Estimation espace disque')

add_table(
    ['Type', 'Taille unitaire', 'Nombre max', 'Espace total'],
    [
        ['Per-user (par admin)', '2-10 Mo', '5 archives', '~50 Mo/admin'],
        ['Global quotidien', '50-200 Mo', '14 archives', '0.7-2.8 Go'],
        ['Global hebdomadaire', '50-200 Mo', '8 archives', '0.4-1.6 Go'],
    ]
)

doc.add_paragraph()
para('Estimation totale pour un EMS avec 5 administrateurs :', bold=True)
add_table(
    ['Composant', 'Calcul', 'Espace'],
    [
        ['Per-user', '5 admins x 50 Mo', '250 Mo'],
        ['Global quotidien', '14 x 200 Mo (max)', '2.8 Go'],
        ['Global hebdomadaire', '8 x 200 Mo (max)', '1.6 Go'],
        ['TOTAL MAXIMUM', '', '4.65 Go'],
    ]
)

doc.add_paragraph()
info_box('Ces estimations sont des maximums theoriques. En pratique, pour un EMS de '
         'taille moyenne (100-150 employes), l\'espace reel sera de l\'ordre de 1-2 Go.')

doc.add_page_break()

# ============================================================
# 9. FAQ
# ============================================================
h1('9. FAQ')

h3('Q : Que se passe-t-il si le disque est plein ?')
para('Le systeme verifie l\'espace disponible avant de creer une sauvegarde. '
     'Si l\'espace est insuffisant, un message d\'erreur est affiche et la sauvegarde '
     'n\'est pas creee. Les sauvegardes existantes ne sont pas affectees.')

h3('Q : Peut-on telecharger une sauvegarde ?')
para('Oui, chaque sauvegarde peut etre telechargee en cliquant sur le bouton de '
     'telechargement dans la liste. Le fichier ZIP peut etre stocke en dehors du '
     'systeme pour archivage supplementaire.')

h3('Q : La restauration est-elle instantanee ?')
para('La restauration per-user est generalement rapide (quelques secondes). '
     'La restauration globale peut prendre plusieurs minutes selon le volume de '
     'donnees. Un indicateur de progression s\'affiche pendant l\'operation.')

h3('Q : Que se passe-t-il si la restauration echoue en cours de route ?')
para('Le systeme utilise des transactions SQL. Si une erreur survient pendant la '
     'restauration, toutes les modifications sont annulees (rollback) et les donnees '
     'restent dans leur etat precedent.')

h3('Q : Qui peut voir les sauvegardes des autres administrateurs ?')
para('Chaque administrateur ne voit que ses propres sauvegardes per-user. '
     'Les sauvegardes globales sont visibles par tous les administrateurs mais '
     'la restauration globale necessite le code d\'acces special.')

h3('Q : Le code d\'acces special est-il le meme pour tous ?')
para('Oui, il y a un seul code d\'acces special pour le systeme. Il est defini '
     'par la direction et partage uniquement avec les personnes autorisees a '
     'effectuer des restaurations globales.')

h3('Q : Les sauvegardes automatiques continuent-elles pendant les vacances ?')
para('Oui, les sauvegardes automatiques (cron) s\'executent independamment de '
     'toute intervention humaine, 365 jours par an.')

# ============================================================
# FIN
# ============================================================
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()
r = p.add_run('--- Fin du document ---')
r.font.color.rgb = GREY_TEXT
r.font.size = Pt(10)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Document genere automatiquement — SpocSpace v1.0')
r.font.color.rgb = GREY_TEXT
r.font.size = Pt(9)

# ── Save ──
output = 'Sauvegarde_Restauration_SpocSpace.docx'
doc.save(output)
print(f'Document genere : {output}')
